import cv2
from io import BytesIO
from docx import Document
from docx.shared import Cm


def generate_flipbook(frames, height, left_margin, sheet_margin, filename_flipbook):
    # Note that the frames come in OpenCVs default BGR format.
    print(f"Create flipbook with {len(frames)} frames ...")
    # We create a document and within this document we create a paragraph. Within this paragraph
    # we create a run, this will allow us to put the images side by side. https://stackoverflow.com/a/57363257
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    # Also, we want to reduce the margins so more images can fit on one page.
    section = doc.sections[0]
    section.top_margin = Cm(sheet_margin)
    section.bottom_margin = Cm(sheet_margin)
    section.left_margin = Cm(sheet_margin)
    section.right_margin = Cm(sheet_margin)
    # Loop over frames and add each frame to the Word document.
    for idx, frame in enumerate(frames):
        # Convert each frame (in RGB format) to an image file in memory (BytesIO).
        image_stream = BytesIO()
        # We want to add a white margin to the left that can be used to connect the paper sheets.
        # We will do this using opencv, because its simpler.
        frame = cv2.copyMakeBorder(
            frame,
            0,  # top
            0,  # bottom
            int(
                frame.shape[0] / height * left_margin
            ),  # left, number of pixels calculated from cm's
            0,  # right
            cv2.BORDER_CONSTANT,
            None,
            [255, 255, 255],
        )
        # Also, add a little border around the sheet so its easier to cut out.
        frame = cv2.copyMakeBorder(
            frame,
            1,  # top
            1,  # bottom
            1,  # left
            1,  # right
            cv2.BORDER_CONSTANT,
            None,
            [200, 200, 200],
        )
        # In the unfortunate case of letting all the cutout sheets fall to the ground it would be
        # nice to still be able to recover the order. So we put little numbers at the bottom left corner.
        cv2.putText(
            frame,
            str(idx + 1),
            (10, frame.shape[0] - 10),
            cv2.FONT_HERSHEY_SIMPLEX,
            3,
            (0, 0, 0),
            2,
        )
        # Encode the frame as png.
        _, encoded_image = cv2.imencode(".png", frame)
        # Write the encoded image to the in-memory BytesIO stream
        image_stream.write(encoded_image.tobytes())
        # Reset pointer to the start of the stream.
        image_stream.seek(0)
        # Add the image to the document.
        run.add_picture(image_stream, height=Cm(height))

    # Save the Word document.
    doc.save(filename_flipbook)
    print(f"Flipbook saved to '{filename_flipbook}'.")
